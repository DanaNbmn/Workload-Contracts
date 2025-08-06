import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import tempfile
import base64

# --- Determine Title Prefix Based on Degree and Gender ---
def determine_title(degree, gender):
    degree = str(degree).lower()
    gender = str(gender).lower()
    if "phd" in degree or "dphil" in degree or "doctorate" in degree:
        return "Dr."
    elif gender == "female":
        return "Ms."
    else:
        return "Mr."

# --- Contract Generator ---
def generate_contracts(df, logo_file):
    output_paths = []
    temp_dir = tempfile.mkdtemp()

    for i in range(len(df)):
        row = df.iloc[i]
        doc = Document()

        if logo_file:
            doc.add_picture(logo_file, width=Inches(1.5))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        title = determine_title(row['Degree'], row['Gender'])
        faculty_name = f"{title} {row['Name']}"

        heading = doc.add_heading('SERVICE AGREEMENT', level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        date_para = doc.add_paragraph()
        date_run = date_para.add_run("This Agreement is made on: [Insert Date]")
        date_run.font.size = Pt(11)
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        doc.add_paragraph("Opening Statement:", style='Heading 2')
        doc.add_paragraph(
            "This Service Agreement is entered into between Abu Dhabi University (hereinafter referred to as the \u201cFirst Party\u201d) and the employee identified below (hereinafter referred to as the \u201cSecond Party\u201d). This Agreement outlines the terms and conditions under which the Second Party will perform academic duties for the specified academic period."
        )

        doc.add_paragraph("\nParties:", style='Heading 2')
        doc.add_paragraph("First Party: Abu Dhabi University")
        doc.add_paragraph(
            f"Second Party:\n\u2022 Name: {faculty_name}\n\u2022 Faculty Type: {row['Faculty Type']}\n\u2022 College/Department: {row['College/Department']}\n\u2022 Faculty ID: {row.get('Faculty ID', 'N/A')}"
        )

        doc.add_paragraph("\nContract Period:", style='Heading 2')
        doc.add_paragraph(f"Academic Year: AY {row['Academic Year']}\nSemester / Term: {row['Semester/Term']}")

        doc.add_paragraph("\nScope of Work:", style='Heading 2')
        scope_points = [
            "Deliver the assigned course(s) in line with the approved schedule and syllabus.",
            "Submit final student grades in accordance with the official academic calendar.",
            "Complete and upload all required course documentation (e.g., course files, assessment materials).",
            "Remain available to address student inquiries, including during any approved post-semester extension period."
        ]
        for point in scope_points:
            doc.add_paragraph(point, style='List Bullet')

        doc.add_paragraph("\nCompensation", style='Heading 2')
        doc.add_paragraph(f"Total Compensation (AED): {row['Compensation (AED)']}")

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Workload Hours'
        hdr_cells[1].text = 'Course Level'
        hdr_cells[2].text = 'Payment Details'
        hdr_cells[3].text = 'Compensation (AED)'
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Workload Hours'])
        row_cells[1].text = row['Course Level']
        row_cells[2].text = row['Payment Details']
        row_cells[3].text = str(row['Compensation (AED)'])

        doc.add_paragraph("\nInstalment Details:", style='Heading 2')
        doc.add_paragraph("\u2022 The total compensation will be paid in equal monthly instalments over the duration of the contract, with each instalment released upon completion of teaching duties and submission of required deliverables (e.g., grades and course files).")
        doc.add_paragraph("\u2022 Instalment payments are conditional upon adherence to Abu Dhabi University\u2019s academic policies and timelines. Any failure to meet contractual obligations may result in payment delays, adjustments, or withholdings.")

        doc.add_paragraph("\nPolicies and Compliance", style='Heading 2')
        compliance_points = [
            "Comply with all applicable Abu Dhabi University policies, procedures, and academic regulations.",
            "Demonstrate professionalism and ethical conduct in all teaching-related activities.",
            "Support institutional quality assurance, accreditation, and review processes as requested."
        ]
        for point in compliance_points:
            doc.add_paragraph(point, style='List Bullet')

        doc.add_paragraph("\nSignatures and Acknowledgement", style='Heading 2')
        doc.add_paragraph("By signing this Agreement, all parties confirm their understanding and acceptance of the terms set forth herein.")

        sign_table = doc.add_table(rows=4, cols=4)
        sign_table.style = 'Table Grid'
        sign_table.cell(0, 0).text = 'Name'
        sign_table.cell(0, 1).text = 'Title'
        sign_table.cell(0, 2).text = 'Signature'
        sign_table.cell(0, 3).text = 'Date'
        sign_table.cell(1, 0).text = row['Dean Name']
        sign_table.cell(1, 1).text = 'Dean / Department Head / Authorized Signatory'
        sign_table.cell(2, 0).text = row['Faculty Name']
        sign_table.cell(2, 1).text = 'Faculty – Second Party'
        sign_table.cell(3, 0).text = row['HR Representative Name']
        sign_table.cell(3, 1).text = 'Representative, Talent Empowerment and Growth Department'

        filename = f"{faculty_name.replace(' ', '_')}_{row['Academic Year'].replace('/', '-')}_{row['Semester/Term'].replace(' ', '_')}.docx"
        file_path = os.path.join(temp_dir, filename)
        doc.save(file_path)
        output_paths.append(file_path)

    return output_paths

# --- Streamlit UI ---
st.set_page_config(layout="centered")
st.title("\U0001F4C4 ADU Faculty Contract Generator")
st.markdown("Upload your Excel file and logo to generate styled, official faculty service agreement contracts.")

uploaded_file = st.file_uploader("Upload Excel File", type=["xlsx"])
logo_file = st.file_uploader("Upload ADU Logo (PNG)", type=["png"])

if uploaded_file is not None:
    try:
        df = pd.read_excel(uploaded_file)
        st.success("Excel file loaded successfully.")

        if st.button("Generate Contracts"):
            with st.spinner("Generating contracts..."):
                contract_paths = generate_contracts(df, logo_file)

            for path in contract_paths:
                with open(path, "rb") as f:
                    b64 = base64.b64encode(f.read()).decode()
                    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{os.path.basename(path)}">Download {os.path.basename(path)}</a>'
                    st.markdown(href, unsafe_allow_html=True)

            st.success("All contracts generated and ready for download.")
    except Exception as e:
        st.error(f"An error occurred while processing the file: {e}")
